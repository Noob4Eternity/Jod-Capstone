"""
User Story Generation and Validation Agents for AI Project Management System
Using LangGraph State Management and Gemini-2.0-flash - FIXED MULTIMODAL VERSION
"""

import json
import os
import re
from typing import TypedDict, List, Dict, Optional, Any, Union
from datetime import datetime
from enum import Enum
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field, field_validator
from langgraph.graph import StateGraph, END

# ==================== CONSTANTS / SCHEMA (API-READY) ====================

USER_STORY_JSON_SCHEMA = r"""{
    "type": "array",
    "items": {
        "type": "object",
        "required": [
            "id",
            "title",
            "description",
            "acceptance_criteria",
            "priority",
            "estimated_points",
            "dependencies",
            "technical_notes"
        ],
        "properties": {
            "id": {"type": "string", "pattern": "^US\\d{3}$"},
            "title": {"type": "string"},
            "description": {"type": "string"},
            "acceptance_criteria": {"type": "array", "minItems": 3, "maxItems": 7, "items": {"type": "string"}},
            "priority": {"type": "string", "enum": ["high", "medium", "low"]},
            "estimated_points": {"type": "integer", "enum": [1,2,3,5,8,13]},
            "dependencies": {"type": "array", "items": {"type": "string", "pattern": "^US\\d{3}$"}},
            "technical_notes": {"type": "string"}
        }
    }
}"""

# ==================== STATE DEFINITION ====================

class ValidationStatus(Enum):
    APPROVED = "approved"
    NEEDS_REVISION = "needs_revision"
    NEEDS_CLARIFICATION = "needs_clarification"

class ProjectManagementState(TypedDict):
    """Shared state for all agents in the workflow"""
    # Input - Enhanced to support documentation
    client_requirements: str  # Backward compatibility
    documentation: Optional[Dict]  # New structured documentation input
    project_id: str
    project_context: Optional[Dict]  # Industry, team size, tech stack, etc.
    
    # Generated artifacts
    parsed_requirements: Optional[Dict]
    user_stories: Optional[List[Dict]]
    previous_user_stories: Optional[List[Dict]]  # For tracking iterations
    tasks: Optional[List[Dict]]  # Generated tasks from user stories
    
    # Validation and Feedback Loop
    validation_status: Optional[str]
    validation_feedback: Optional[List[str]]
    validation_score: Optional[float]
    story_issues: Optional[Dict[str, List[str]]]  # Issues per story ID
    detailed_feedback: Optional[Dict]  # Rich feedback from validation agent
    improvement_instructions: Optional[List[str]]  # Specific instructions for next iteration
    
    # Metadata
    current_phase: str
    iteration_count: int
    max_iterations: int
    last_error: Optional[str]
    processing_time: Optional[Dict[str, float]]
    timestamp: str

# ==================== UTILITY FUNCTIONS FOR TYPE SAFETY ====================

def safe_string_extract(obj: Any) -> str:
    """Safely extract string from various object types"""
    if isinstance(obj, str):
        return obj
    elif isinstance(obj, dict):
        # Try common string fields in dictionaries
        for key in ['text', 'content', 'description', 'title', 'value']:
            if key in obj and isinstance(obj[key], str):
                return obj[key]
        # Fall back to string representation
        return str(obj)
    elif isinstance(obj, (list, tuple)):
        # Join list elements if they're strings
        string_items = [safe_string_extract(item) for item in obj if item]
        return ' '.join(string_items)
    elif obj is None:
        return ""
    else:
        return str(obj)

def safe_list_extract(obj: Any) -> List[str]:
    """Safely extract list of strings from various object types"""
    if isinstance(obj, list):
        return [safe_string_extract(item) for item in obj if item]
    elif isinstance(obj, str):
        # Split string by common delimiters
        if ',' in obj:
            return [item.strip() for item in obj.split(',') if item.strip()]
        elif ';' in obj:
            return [item.strip() for item in obj.split(';') if item.strip()]
        else:
            return [obj] if obj.strip() else []
    elif obj is None:
        return []
    else:
        return [safe_string_extract(obj)]

def normalize_analysis_data(analysis: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize analysis data to ensure consistent types"""
    normalized = {}
    
    # Define expected fields and their types
    field_configs = {
        'core_features': (list, []),
        'stakeholders': (list, ['user']),
        'technical_constraints': (list, []),
        'business_goals': (list, []),
        'conflicts': (list, []),
        'gaps': (list, [])
    }
    
    for field, (expected_type, default_value) in field_configs.items():
        raw_value = analysis.get(field, default_value)
        
        if expected_type == list:
            normalized[field] = safe_list_extract(raw_value)
        else:
            normalized[field] = default_value
    
    return normalized

# ==================== ENHANCED USER STORY GENERATION AGENT ====================

class MultimodalUserStoryGenerationAgent:
    """Enhanced agent for generating user stories from multimodal inputs (text + PDF)"""
    
    def __init__(self, gemini_api_key: str = None, temperature: float = 0.3):
        api_key = gemini_api_key or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY must be provided either as parameter or environment variable")
        
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-pro",
            temperature=0.8,
            google_api_key=api_key
        )
        
        # Enhanced multimodal story generation prompt
        self.multimodal_story_prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """You are an expert Product Manager + Agile BA generating HIGH-QUALITY user stories from MULTIMODAL requirements.

You receive requirements from multiple sources with different priorities:
1. PRIMARY REQUIREMENTS (user text input) - HIGHEST PRIORITY
2. SUPPORTING DOCUMENTATION (PDF content) - CONTEXT & REFERENCE
3. PROJECT CONTEXT - TECHNICAL & BUSINESS CONSTRAINTS

MULTIMODAL PROCESSING RULES:
- PRIMARY requirements take precedence over document content
- Use documents to add context, clarify ambiguities, and fill gaps
- Identify and resolve conflicts between sources (favor primary requirements)
- Extract stakeholders, constraints, and success criteria from all sources
- Generate stories that satisfy primary requirements while incorporating relevant document insights

OUTPUT CONTRACT: {json_schema}

QUALITY STANDARDS:
- Format: "As a [persona], I want [capability] so that [benefit]"
- Acceptance Criteria: 3-7 measurable, testable criteria per story
- Dependencies: Valid story IDs only, no circular dependencies
- Technical Notes: Actionable implementation guidance
- Coverage: Address ALL primary requirements and relevant document features

SCORING ALIGNMENT (100 points):
Format(15) + Completeness(20) + Requirements Coverage(25) + Source Integration(10) + NFR Coverage(10) + Dependencies(10) + Acceptance Criteria Quality(10)

Return ONLY the JSON array of story objects."""
            ),
            (
                "human",
                """MULTIMODAL REQUIREMENTS INPUT:

=== PRIMARY REQUIREMENTS (User Input - HIGHEST PRIORITY) ===
{primary_requirements}

=== SUPPORTING DOCUMENTATION (PDF Content - Context & Reference) ===
{document_content}

=== SOURCE ANALYSIS ===
{source_analysis}

=== CONFLICT RESOLUTION ===
{conflict_resolution}

=== PROJECT CONTEXT ===
{project_context}

=== VALIDATION FEEDBACK (if applicable) ===
{feedback_section}

INSTRUCTIONS:
1. Prioritize PRIMARY REQUIREMENTS as the main specification
2. Use SUPPORTING DOCUMENTATION to enhance, clarify, and provide context
3. Resolve any conflicts between sources (favor primary requirements)
4. Extract personas from all sources (stakeholders, roles mentioned)
5. Include security/performance/scalability if implied in any source
6. Generate comprehensive technical notes incorporating insights from all sources
7. Ensure all primary requirements are fully covered
{iteration_instructions}

OUTPUT: JSON array ONLY. No surrounding text.
{feedback_focus}"""
            ),
        ])
        
        # Content analysis prompt for multimodal processing
        self.content_analysis_prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """Analyze multimodal requirements input to extract key information for story generation.

IMPORTANT: Return all data as simple strings and arrays of strings only. No nested objects.

Extract and categorize:
1. Core Features: Main functionalities required (return as array of strings)
2. Stakeholders: All user types, roles mentioned across sources (return as array of strings)
3. Technical Constraints: Technology, performance, security requirements (return as array of strings)
4. Business Goals: Objectives, success criteria, KPIs (return as array of strings)
5. Conflicts: Any contradictions between primary and document sources (return as array of strings)
6. Gaps: Missing information that needs clarification (return as array of strings)

Return structured analysis as JSON with arrays of strings only."""
            ),
            (
                "human",
                """PRIMARY REQUIREMENTS:
{primary_text}

DOCUMENT CONTENT:
{document_text}

Analyze and return JSON with: core_features, stakeholders, technical_constraints, business_goals, conflicts, gaps
All fields must be arrays of strings. No nested objects or complex data structures."""
            ),
        ])
        
        self.parser = JsonOutputParser()
    
    def _parse_multimodal_documentation(self, state: ProjectManagementState) -> Dict[str, Any]:
        """
        Parse and organize multimodal content by source type and priority.
        Returns structured analysis for intelligent story generation.
        """
        print(f"[DEBUG] === PARSING MULTIMODAL INPUT ===")
        print(f"[DEBUG] State keys: {list(state.keys())}")
        print(f"[DEBUG] client_requirements length: {len(state.get('client_requirements', ''))}")
        print(f"[DEBUG] documentation exists: {bool(state.get('documentation'))}")
        
        # Extract different content sources
        primary_requirements = ""
        document_content = ""
        source_metadata = {
            "has_primary_text": False,
            "has_documents": False,
            "document_count": 0,
            "content_distribution": {}
        }
        
        # Check for structured documentation (from unified API)
        if state.get("documentation"):
            doc = state["documentation"]
            full_content = doc.get("content", "")
            
            print(f"[DEBUG] documentation content length: {len(full_content)}")
            print(f"[DEBUG] content preview: {full_content[:300]}...")
            
            # Parse structured multimodal content
            if "=== PROJECT REQUIREMENTS (TEXT) ===" in full_content:
                print("[DEBUG] ✓ Found multimodal structure")
                # Split multimodal content
                parts = full_content.split("=== PROJECT REQUIREMENTS (TEXT) ===")
                if len(parts) > 1:
                    # Extract primary requirements
                    primary_section = parts[1].split("=== DOCUMENT:")[0]
                    primary_requirements = primary_section.strip()
                    source_metadata["has_primary_text"] = True
                    
                    # Extract document sections
                    document_sections = []
                    doc_parts = full_content.split("=== DOCUMENT:")
                    for i, part in enumerate(doc_parts[1:], 1):
                        document_sections.append(f"Document {i}: {part.strip()}")
                        source_metadata["document_count"] += 1
                    
                    if document_sections:
                        document_content = "\n\n".join(document_sections)
                        source_metadata["has_documents"] = True
            else:
                print("[DEBUG] ✗ No multimodal structure found - treating as document-only")
                # Single source content (likely PDF only) - treat as primary requirements
                primary_requirements = full_content
                source_metadata["has_primary_text"] = True
                source_metadata["document_count"] = 1
        
        # Fallback to client_requirements (backward compatibility)
        if not primary_requirements and not document_content:
            primary_requirements = state.get("client_requirements", "")
            source_metadata["has_primary_text"] = bool(primary_requirements)
            print(f"[DEBUG] Using client_requirements fallback: {len(primary_requirements)} chars")
        
        # Calculate content distribution
        total_length = len(primary_requirements) + len(document_content)
        if total_length > 0:
            source_metadata["content_distribution"] = {
                "primary_percentage": len(primary_requirements) / total_length * 100,
                "document_percentage": len(document_content) / total_length * 100
            }
        
        print(f"[DEBUG] Final primary_requirements: {len(primary_requirements)} chars")
        print(f"[DEBUG] Final document_content: {len(document_content)} chars")
        print(f"[DEBUG] Source metadata: {source_metadata}")
        
        return {
            "primary_requirements": primary_requirements,
            "document_content": document_content,
            "source_metadata": source_metadata,
            "project_context": state.get("project_context", {})
        }
    
    def _analyze_multimodal_content(self, primary_text: str, document_text: str) -> Dict[str, Any]:
        """
        Analyze multimodal content to extract structured insights for story generation.
        Returns analysis of features, stakeholders, conflicts, etc.
        """
        if not primary_text and not document_text:
            return normalize_analysis_data({
                "core_features": [],
                "stakeholders": ["user", "administrator"],
                "technical_constraints": [],
                "business_goals": [],
                "conflicts": [],
                "gaps": ["No requirements provided"]
            })
        
        try:
            chain = self.content_analysis_prompt | self.llm | self.parser
            analysis = chain.invoke({
                "primary_text": primary_text or "No primary text provided",
                "document_text": document_text or "No document content provided"
            })
            
            # Normalize the analysis to ensure consistent types
            normalized_analysis = normalize_analysis_data(analysis)
            
            print(f"[CONTENT_ANALYSIS] Extracted {len(normalized_analysis.get('core_features', []))} features")
            print(f"[CONTENT_ANALYSIS] Identified {len(normalized_analysis.get('stakeholders', []))} stakeholders")
            
            return normalized_analysis
            
        except Exception as e:
            print(f"[CONTENT_ANALYSIS_ERROR] {e}")
            # Return basic analysis on failure
            return normalize_analysis_data({
                "core_features": ["Core functionality from requirements"],
                "stakeholders": ["user", "administrator"],
                "technical_constraints": [],
                "business_goals": [],
                "conflicts": [f"Analysis failed: {str(e)}"],
                "gaps": ["Manual review recommended"]
            })
    
    def _resolve_source_conflicts(self, analysis: Dict[str, Any], primary_text: str, document_text: str) -> str:
        """
        Generate conflict resolution guidance for the LLM prompt.
        """
        conflicts = analysis.get("conflicts", [])
        
        if not conflicts:
            return "No conflicts detected between sources. Integrate all requirements harmoniously."
        
        resolution_guidance = ["CONFLICTS DETECTED - RESOLUTION STRATEGY:"]
        
        for i, conflict in enumerate(conflicts, 1):
            conflict_str = safe_string_extract(conflict)
            resolution_guidance.append(f"{i}. CONFLICT: {conflict_str}")
            resolution_guidance.append(f"   RESOLUTION: Prioritize primary requirements, use documents for context only")
        
        resolution_guidance.append("\nGENERAL RULE: When in doubt, favor explicit primary requirements over document implications.")
        
        return "\n".join(resolution_guidance)
    
    def _create_source_analysis_summary(self, analysis: Dict[str, Any], source_metadata: Dict[str, Any]) -> str:
        """
        Create a summary of source analysis for the LLM prompt.
        """
        summary_parts = []
        
        try:
            # Source composition
            if source_metadata.get("has_primary_text") and source_metadata.get("has_documents"):
                dist = source_metadata.get("content_distribution", {})
                summary_parts.append(f"MULTIMODAL INPUT: {dist.get('primary_percentage', 0):.0f}% primary text, {dist.get('document_percentage', 0):.0f}% documents")
            elif source_metadata.get("has_primary_text"):
                summary_parts.append("INPUT TYPE: Primary text requirements only")
            elif source_metadata.get("has_documents"):
                summary_parts.append(f"INPUT TYPE: Document-based requirements ({source_metadata.get('document_count', 1)} documents)")
            
            # Key extracted elements
            core_features = analysis.get('core_features', [])
            summary_parts.append(f"CORE FEATURES IDENTIFIED: {len(core_features)}")
            
            # Safe handling of stakeholders
            stakeholders = analysis.get('stakeholders', ['user'])
            if stakeholders:
                stakeholder_strings = [safe_string_extract(s) for s in stakeholders[:5] if s]
                stakeholder_strings = [s for s in stakeholder_strings if s.strip()]  # Filter empty
                if stakeholder_strings:
                    summary_parts.append(f"STAKEHOLDERS IDENTIFIED: {', '.join(stakeholder_strings)}")
            
            # Safe handling of technical constraints
            tech_constraints = analysis.get("technical_constraints", [])
            if tech_constraints:
                constraint_strings = [safe_string_extract(c) for c in tech_constraints[:3] if c]
                constraint_strings = [c for c in constraint_strings if c.strip()]  # Filter empty
                if constraint_strings:
                    summary_parts.append(f"TECHNICAL CONSTRAINTS: {', '.join(constraint_strings)}")
            
            # Safe handling of business goals
            business_goals = analysis.get("business_goals", [])
            if business_goals:
                goal_strings = [safe_string_extract(g) for g in business_goals[:3] if g]
                goal_strings = [g for g in goal_strings if g.strip()]  # Filter empty
                if goal_strings:
                    summary_parts.append(f"BUSINESS GOALS: {', '.join(goal_strings)}")
            
            # Safe handling of gaps
            gaps = analysis.get("gaps", [])
            if gaps:
                gap_strings = [safe_string_extract(g) for g in gaps[:2] if g]
                gap_strings = [g for g in gap_strings if g.strip()]  # Filter empty
                if gap_strings:
                    summary_parts.append(f"GAPS IDENTIFIED: {', '.join(gap_strings)}")
            
        except Exception as e:
            print(f"[SOURCE_ANALYSIS_ERROR] Error creating summary: {e}")
            summary_parts = [
                "INPUT TYPE: Requirements provided",
                "CORE FEATURES IDENTIFIED: Analysis in progress",
                "STAKEHOLDERS IDENTIFIED: user, administrator"
            ]
        
        return "\n".join(summary_parts)
    
    def _format_multimodal_feedback(self, state: ProjectManagementState) -> tuple:
        """
        Format validation feedback specifically for multimodal iteration improvements.
        """
        iteration_count = state.get("iteration_count", 0)
        
        if iteration_count == 0:
            return "", "", ""
        
        feedback_section = "\n=== MULTIMODAL VALIDATION FEEDBACK ===\n"
        iteration_instructions = "\n6. **CRITICAL**: Address multimodal feedback and improve source integration"
        feedback_focus = "\n6. **Improve multimodal processing and source coverage**"
        
        detailed_feedback = state.get("detailed_feedback", {})
        
        if detailed_feedback:
            feedback_section += f"**Validation Score**: {state.get('validation_score', 0):.1f}/100\n\n"
            
            # Source coverage issues
            if detailed_feedback.get("missing_requirements"):
                feedback_section += "**MISSING REQUIREMENTS (check all sources)**:\n"
                for req in detailed_feedback["missing_requirements"]:
                    feedback_section += f"- {req}\n"
                feedback_section += "\n"
            
            # Multimodal integration feedback
            if detailed_feedback.get("recommendations"):
                feedback_section += "**MULTIMODAL INTEGRATION IMPROVEMENTS**:\n"
                for rec in detailed_feedback["recommendations"]:
                    feedback_section += f"- {rec}\n"
                feedback_section += "\n"
            
            # Critical issues
            if detailed_feedback.get("critical_issues"):
                feedback_section += "**CRITICAL ISSUES TO FIX**:\n"
                for issue in detailed_feedback["critical_issues"]:
                    feedback_section += f"- {issue}\n"
                feedback_section += "\n"
        
        # Previous iteration reference
        previous_stories = state.get("previous_user_stories", [])
        if previous_stories:
            feedback_section += f"**PREVIOUS ITERATION**: {len(previous_stories)} stories generated - improve based on feedback\n"
        
        return feedback_section, iteration_instructions, feedback_focus
    
    def generate_stories(self, state: ProjectManagementState) -> ProjectManagementState:
        """
        Enhanced story generation method with full multimodal support.
        """
        start_time = datetime.now()
        
        try:
            # Store previous stories
            current_stories = state.get("user_stories", [])
            if current_stories:
                state["previous_user_stories"] = current_stories
            
            # Parse multimodal content
            multimodal_data = self._parse_multimodal_documentation(state)
            primary_requirements = multimodal_data["primary_requirements"]
            document_content = multimodal_data["document_content"]
            source_metadata = multimodal_data["source_metadata"]
            project_context = multimodal_data["project_context"]
            
            # Analyze content across sources
            content_analysis = self._analyze_multimodal_content(primary_requirements, document_content)
            
            # Create source analysis summary
            source_analysis = self._create_source_analysis_summary(content_analysis, source_metadata)
            
            # Resolve conflicts between sources
            conflict_resolution = self._resolve_source_conflicts(content_analysis, primary_requirements, document_content)
            
            # Format validation feedback for iteration
            feedback_section, iteration_instructions, feedback_focus = self._format_multimodal_feedback(state)
            
            # Log multimodal processing info
            iteration_count = state.get("iteration_count", 0)
            print(f"[MULTIMODAL_GEN] Iteration {iteration_count + 1}")
            print(f"[MULTIMODAL_GEN] Primary text: {len(primary_requirements)} chars")
            print(f"[MULTIMODAL_GEN] Document content: {len(document_content)} chars")
            print(f"[MULTIMODAL_GEN] Features identified: {len(content_analysis.get('core_features', []))}")
            
            # Safe stakeholder display
            stakeholders = content_analysis.get('stakeholders', [])
            if stakeholders:
                stakeholder_preview = [safe_string_extract(s) for s in stakeholders[:3]]
                stakeholder_preview = [s for s in stakeholder_preview if s.strip()]
                print(f"[MULTIMODAL_GEN] Stakeholders: {', '.join(stakeholder_preview)}")
            
            if content_analysis.get('conflicts'):
                print(f"[MULTIMODAL_GEN] Conflicts detected: {len(content_analysis['conflicts'])}")
            
            # Generate stories using enhanced multimodal prompt
            chain = self.multimodal_story_prompt | self.llm | self.parser
            
            raw_stories = chain.invoke({
                "primary_requirements": primary_requirements or "No primary requirements provided",
                "document_content": document_content or "No supporting documentation provided",
                "source_analysis": source_analysis,
                "conflict_resolution": conflict_resolution,
                "project_context": json.dumps(project_context) if project_context else "No specific context",
                "feedback_section": feedback_section,
                "iteration_instructions": iteration_instructions,
                "feedback_focus": feedback_focus,
                "json_schema": USER_STORY_JSON_SCHEMA
            })
            
            # Ensure we have a list
            if not isinstance(raw_stories, list):
                raw_stories = [raw_stories] if raw_stories else []
            
            # Validate and enhance each story with multimodal insights
            validated_stories = []
            for idx, story in enumerate(raw_stories):
                story_id = story.get("id", f"US{idx+1:03d}")
                
                # Enhanced technical notes with multimodal insights
                technical_notes = story.get("technical_notes", "")
                tech_constraints = content_analysis.get("technical_constraints", [])
                if tech_constraints:
                    constraint_strings = [safe_string_extract(c) for c in tech_constraints[:2] if c]
                    constraint_strings = [c for c in constraint_strings if c.strip()]
                    if constraint_strings:
                        technical_notes += f" Consider: {', '.join(constraint_strings)}"
                
                validated_story = {
                    "id": story_id,
                    "title": story.get("title", ""),
                    "description": story.get("description", ""),
                    "acceptance_criteria": story.get("acceptance_criteria", []),
                    "priority": story.get("priority", "medium"),
                    "estimated_points": story.get("estimated_points", 3),
                    "dependencies": story.get("dependencies", []),
                    "technical_notes": technical_notes.strip(),
                    "source_traceability": {
                        "primary_coverage": bool(primary_requirements and any(
                            word in story.get("title", "").lower() + story.get("description", "").lower()
                            for word in primary_requirements.lower().split()[:10]
                        )),
                        "document_coverage": bool(document_content and any(
                            word in story.get("title", "").lower() + story.get("description", "").lower()
                            for word in document_content.lower().split()[:10]
                        ))
                    }
                }
                
                # Fix story format if needed
                if validated_story["title"] and not self._is_valid_story_format(validated_story["title"]):
                    validated_story["title"] = self._fix_story_format(validated_story["title"])
                
                # Ensure minimum acceptance criteria
                if len(validated_story["acceptance_criteria"]) < 3:
                    while len(validated_story["acceptance_criteria"]) < 3:
                        base_criteria = validated_story["acceptance_criteria"][-1] if validated_story["acceptance_criteria"] else "System responds successfully"
                        validated_story["acceptance_criteria"].append(f"Enhanced: {base_criteria}")
                
                validated_stories.append(validated_story)
            
            # Ensure sequential IDs
            for i, story in enumerate(validated_stories):
                story["id"] = f"US{i+1:03d}"
            
            # Enhanced coverage check with multimodal awareness
            validated_stories = self._ensure_multimodal_coverage(
                primary_requirements, document_content, content_analysis, validated_stories
            )
            
            # Update state with multimodal metadata
            state["user_stories"] = validated_stories
            state["current_phase"] = "story_validation"
            state["multimodal_metadata"] = {
                "source_analysis": content_analysis,
                "source_distribution": source_metadata.get("content_distribution", {}),
                "conflicts_resolved": len(content_analysis.get("conflicts", [])),
                "stories_with_primary_coverage": sum(1 for s in validated_stories if s.get("source_traceability", {}).get("primary_coverage", False)),
                "stories_with_document_coverage": sum(1 for s in validated_stories if s.get("source_traceability", {}).get("document_coverage", False))
            }
            
            processing_time = (datetime.now() - start_time).total_seconds()
            if "processing_time" not in state:
                state["processing_time"] = {}
            state["processing_time"]["multimodal_story_generation"] = processing_time
            
            print(f"[MULTIMODAL_GEN] Generated {len(validated_stories)} stories in {processing_time:.1f}s")
            
        except Exception as e:
            state["last_error"] = f"Multimodal story generation failed: {str(e)}"
            state["user_stories"] = []
            state["current_phase"] = "error"
            print(f"[MULTIMODAL_GEN_ERROR] {str(e)}")
            import traceback
            traceback.print_exc()
        
        return state
    
    def _ensure_multimodal_coverage(self, primary_text: str, document_text: str, analysis: Dict[str, Any], stories: List[Dict]) -> List[Dict]:
        """
        Enhanced coverage check that considers both primary requirements and document content.
        """
        all_content = f"{primary_text} {document_text}".lower()
        story_content = " ".join([
            f"{s.get('title', '')} {s.get('description', '')} {' '.join(s.get('acceptance_criteria', []))}"
            for s in stories
        ]).lower()
        
        missing_elements = []
        
        # Check coverage of identified core features (with safe string extraction)
        core_features = analysis.get("core_features", [])
        for feature in core_features[:5]:  # Check top 5 features
            feature_str = safe_string_extract(feature).lower()
            if feature_str and feature_str not in story_content:
                missing_elements.append(f"core_feature_{len(missing_elements)}")
        
        # Check technical constraints coverage
        technical_constraints = analysis.get("technical_constraints", [])
        for constraint in technical_constraints:
            constraint_str = safe_string_extract(constraint).lower()
            constraint_keywords = ["security", "performance", "scalability", "authentication"]
            if any(keyword in constraint_str for keyword in constraint_keywords):
                if not any(keyword in story_content for keyword in constraint_keywords):
                    missing_elements.append("technical_constraints")
                    break
        
        # Generate missing stories if needed
        if missing_elements:
            additional_stories = self._generate_multimodal_gap_stories(missing_elements, analysis, len(stories))
            stories.extend(additional_stories)
        
        return stories
    
    def _generate_multimodal_gap_stories(self, missing_elements: List[str], analysis: Dict[str, Any], current_count: int) -> List[Dict]:
        """
        Generate stories to fill gaps identified in multimodal analysis.
        """
        gap_stories = []
        
        # Templates enhanced with multimodal insights
        stakeholders = analysis.get("stakeholders", ["user", "administrator"])
        primary_stakeholder = "user"
        if stakeholders:
            primary_stakeholder_raw = stakeholders[0] if stakeholders else "user"
            primary_stakeholder = safe_string_extract(primary_stakeholder_raw)
            if not primary_stakeholder.strip():
                primary_stakeholder = "user"
        
        for idx, element in enumerate(missing_elements[:3]):  # Limit to 3 additional stories
            if element.startswith("core_feature_"):
                story = {
                    "id": f"US{current_count + idx + 1:03d}",
                    "title": f"As a {primary_stakeholder}, I want to access core system functionality so that I can complete essential tasks",
                    "description": "Implement core feature identified in requirements analysis but not covered by existing stories",
                    "acceptance_criteria": [
                        "Core functionality is accessible through the UI",
                        "Feature works according to requirements specification",
                        "Error handling is implemented for edge cases",
                        "Feature integrates properly with existing system"
                    ],
                    "priority": "high",
                    "estimated_points": 5,
                    "dependencies": [],
                    "technical_notes": f"Implement missing core feature identified in multimodal analysis. Reference both primary requirements and supporting documentation."
                }
            elif element == "technical_constraints":
                # Safe extraction of technical constraints
                tech_constraints = analysis.get('technical_constraints', [])
                constraint_strings = [safe_string_extract(c) for c in tech_constraints[:3] if c]
                constraint_strings = [c for c in constraint_strings if c.strip()]
                constraints_text = ', '.join(constraint_strings) if constraint_strings else "various technical requirements"
                
                story = {
                    "id": f"US{current_count + idx + 1:03d}",
                    "title": f"As a system administrator, I want robust technical infrastructure so that the system meets all constraints",
                    "description": "Implement technical requirements and constraints identified across multiple requirement sources",
                    "acceptance_criteria": [
                        "System meets performance requirements",
                        "Security constraints are properly implemented",
                        "Scalability requirements are addressed",
                        "Technical documentation is complete"
                    ],
                    "priority": "high",
                    "estimated_points": 8,
                    "dependencies": [],
                    "technical_notes": f"Address technical constraints from multimodal analysis: {constraints_text}"
                }
            else:
                # Generic gap story
                story = {
                    "id": f"US{current_count + idx + 1:03d}",
                    "title": f"As a {primary_stakeholder}, I want complete system coverage so that all requirements are addressed",
                    "description": "Address requirements gap identified in multimodal analysis",
                    "acceptance_criteria": [
                        "Gap in requirements is properly addressed",
                        "Implementation aligns with both primary and document requirements",
                        "Functionality is tested and validated"
                    ],
                    "priority": "medium",
                    "estimated_points": 3,
                    "dependencies": [],
                    "technical_notes": "Fill requirements gap identified through multimodal content analysis"
                }
            
            gap_stories.append(story)
        
        return gap_stories
    
    def _is_valid_story_format(self, title: str) -> bool:
        """Check if story follows the standard format"""
        pattern = r"^As a .+, I want .+ so that .+"
        return bool(re.match(pattern, title, re.IGNORECASE))
    
    def _fix_story_format(self, title: str) -> str:
        """Attempt to fix story format"""
        if "want" in title.lower():
            return title
        return f"As a user, I want {title} so that I can complete my tasks"

# ==================== ENHANCED VALIDATION AGENT (FIXED FOR ACCURATE SCORING) ====================

class EnhancedUserStoryValidationAgent:
    """
    Enhanced validation agent that validates against the full requirements context
    while maintaining stability fixes for the Pro model.
    """
    
    def __init__(self, gemini_api_key: str = None, temperature: float = 0.2):
        api_key = gemini_api_key or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY must be provided either as parameter or environment variable")
        
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-pro",
            temperature=temperature,
            google_api_key=api_key
        )
        
        # CHANGED: The prompt now accepts the full, original requirements again.
        # This is critical for the agent to accurately score source coverage.
        self.validation_prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """You are an expert QA / Agile validator for MULTIMODAL user story generation.

Your task is to validate the `GENERATED USER STORIES` against the `ORIGINAL REQUIREMENTS SOURCES`. You must ensure every part of the original requirements is covered.

Assess user stories using this weighted rubric (total 100):
1.  **Requirements & NFR Coverage (25 points):** How well do the stories cover ALL primary and document requirements?
2.  **Multimodal Source Integration (15 points):** Do the stories correctly synthesize details from both text and documents?
3.  **Completeness & Acceptance Criteria Quality (20 points):** Are the stories and their ACs clear, testable, and complete?
4.  **Format & Required Fields (15 points):** Do all stories follow the required schema?
5.  **Dependencies & Consistency (10 points):** Are dependencies logical? Are there contradictions?
6.  **Sizing & Priority Logic (10 points):** Are story points and priorities reasonable?
7.  **Technical Notes Quality (5 points):** Are technical notes helpful and relevant?

Return a detailed JSON analysis. The `validation_score` and `multimodal_analysis` scores are crucial."""
            ),
            (
                "human",
                """ORIGINAL REQUIREMENTS SOURCES:

=== PRIMARY REQUIREMENTS ===
{primary_requirements}

=== SUPPORTING DOCUMENTATION ===  
{document_content}

=== MULTIMODAL METADATA (for context) ===
{multimodal_metadata}

=== GENERATED USER STORIES (for validation) ===
{stories}

INSTRUCTIONS:
Validate the generated stories against ALL original sources. Return ONLY a JSON object with the specified keys:
- overall_valid: boolean
- validation_score: float (0-100)
- missing_requirements: string[] (List specific features from any source that were missed)
- story_issues: object mapping story ID to string[] of issues
- recommendations: string[]
- critical_issues: string[]
- warnings: string[]
- multimodal_analysis: object with source_coverage_score, integration_quality, conflict_resolution_score
"""
            ),
        ])
        
        self.parser = JsonOutputParser()

    def _safe_to_float(self, value: Any, default: float = 50.0) -> float:
        """Safely converts a value to a float, returning a default if conversion fails."""
        if value is None: return default
        try: return float(value)
        except (ValueError, TypeError): return default

    # RE-ADDED: This helper function is now needed again to get the full text.
    def _extract_requirements_from_state(self, state: ProjectManagementState) -> Dict[str, str]:
        """Extracts full requirements from state, handling both legacy and multimodal inputs."""
        primary_requirements = ""
        document_content = ""
        
        if state.get("documentation"):
            doc = state["documentation"]
            full_content = doc.get("content", "")
            
            if "=== PROJECT REQUIREMENTS (TEXT) ===" in full_content:
                parts = full_content.split("=== PROJECT REQUIREMENTS (TEXT) ===")
                if len(parts) > 1:
                    primary_section = parts[1].split("=== DOCUMENT:")[0]
                    primary_requirements = primary_section.strip()
                    
                    doc_parts = full_content.split("=== DOCUMENT:")
                    if len(doc_parts) > 1:
                        document_sections = [part.strip() for part in doc_parts[1:]]
                        document_content = "\n\n".join([f"Document: {section}" for section in document_sections])
            else:
                primary_requirements = full_content
        
        if not primary_requirements and not document_content:
            primary_requirements = state.get("client_requirements", "")
        
        return {
            "primary_requirements": primary_requirements,
            "document_content": document_content
        }

    def validate_stories(self, state: ProjectManagementState) -> ProjectManagementState:
        """Enhanced validation method with full context for accurate scoring."""
        start_time = datetime.now()
        try:
            stories = state.get("user_stories", [])
            print(f"[VALIDATION] Validating {len(stories)} stories with full context prompt...")

            if not stories:
                # Same as before
                state["validation_status"] = ValidationStatus.NEEDS_REVISION.value
                state["validation_feedback"] = ["No user stories generated to validate"]
                state["validation_score"] = 0.0
                return state

            # RE-ADDED: Get the full requirements text instead of just summaries.
            requirements_data = self._extract_requirements_from_state(state)
            multimodal_metadata = state.get("multimodal_metadata", {})

            print(f"[VALIDATION] Calling gemini-2.5-pro with {len(requirements_data['primary_requirements'])} chars of primary requirements...")

            try:
                chain = self.validation_prompt | self.llm | self.parser
                
                # CHANGED: Pass the full requirements and stories to the prompt.
                semantic_validation = chain.invoke({
                    "primary_requirements": requirements_data["primary_requirements"] or "No primary requirements provided",
                    "document_content": requirements_data["document_content"] or "No supporting documentation provided",
                    "multimodal_metadata": json.dumps(multimodal_metadata),
                    "stories": json.dumps(stories, indent=2)
                })
                
            except Exception as validation_error:
                # Same fallback logic
                print(f"[VALIDATION_ERROR] LLM validation failed: {validation_error}")
                semantic_validation = { "validation_score": 50.0, "overall_valid": False, "critical_issues": [f"Semantic validation failed: {str(validation_error)}"], "recommendations": ["Manual review required."], "multimodal_analysis": {}}
            
            # This logic remains the same, but now it will have accurate scores to work with.
            base_score = self._safe_to_float(semantic_validation.get("validation_score"), default=70.0)
            multimodal_analysis = semantic_validation.get("multimodal_analysis", {})

            if multimodal_analysis:
                multimodal_bonus = (
                    self._safe_to_float(multimodal_analysis.get("source_coverage_score")) +  
                    self._safe_to_float(multimodal_analysis.get("integration_quality")) +  
                    self._safe_to_float(multimodal_analysis.get("conflict_resolution_score"))
                ) / 3
                if multimodal_bonus > 70: base_score += min(10, (multimodal_bonus - 70) / 3)
            
            validation_score = max(0, min(100, base_score))
            
            # The rest of the state update logic is the same...
            if validation_score >= 80:
                validation_status = ValidationStatus.APPROVED.value
                next_phase = "generate_tasks" 
            elif validation_score >= 60:
                validation_status = ValidationStatus.NEEDS_REVISION.value
                next_phase = "story_generation"
            else:
                validation_status = ValidationStatus.NEEDS_CLARIFICATION.value
                next_phase = "requirement_parsing"

            state["validation_status"] = validation_status
            state["validation_score"] = validation_score
            state["detailed_feedback"] = semantic_validation
            state["story_issues"] = semantic_validation.get("story_issues", {})
            state["current_phase"] = next_phase if validation_status == ValidationStatus.APPROVED.value else "story_generation"

            if validation_status != ValidationStatus.APPROVED.value:
                state["iteration_count"] = state.get("iteration_count", 0) + 1

            print(f"[VALIDATION] Score: {validation_score:.1f}/100, Status: {validation_status}")
            if multimodal_analysis:
                print(f"[VALIDATION] Source Coverage: {self._safe_to_float(multimodal_analysis.get('source_coverage_score')):.1f}/100, Integration Quality: {self._safe_to_float(multimodal_analysis.get('integration_quality')):.1f}/100")

        except Exception as e:
            # Same fatal error handling
            state["last_error"] = f"Enhanced validation failed unexpectedly: {str(e)}"
            state["validation_status"] = ValidationStatus.NEEDS_REVISION.value
            state["current_phase"] = "error"
            print(f"FATAL VALIDATION ERROR: {e}")
            import traceback
            traceback.print_exc()
        
        return state
        
# ==================== TASK GENERATION AGENT ====================

class TaskGenerationAgent:
    """Agent responsible for generating tasks from validated user stories"""
    
    def __init__(self, gemini_api_key: str = None, temperature: float = 0.3):
        api_key = gemini_api_key or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY must be provided either as parameter or environment variable")
        
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-pro",
            temperature=0.4,
            google_api_key=api_key
        )
        
        # Optimized batch processing prompt for multiple stories
        self.batch_task_prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """You are a Technical Lead decomposing user stories into development tasks.

For each story, generate 3-5 tasks:
- Backend task (API/database)
- Frontend task (UI/components) 
- Testing task (unit/integration tests)
- Optional: DevOps or documentation task

Each task: 4-16 hours, atomic, includes story_id.
Each task must include a concise, one-sentence description explaining its purpose.

Return JSON array: [{{
  "story_id": "US001",
  "title": "Task name",
  "description": "Create REST API endpoints for user registration, login, and logout.",
  "category": "backend|frontend|testing|devops",
  "estimated_hours": 8,
  "dependencies": []
}}]"""
            ),
            (
                "human",
                """Stories:
{stories_formatted}

Tech: {tech_stack}

Generate 3-5 tasks per story. Include a unique title and a concise description for each task. Return JSON array only."""
            ),
        ])
        
        self.parser = JsonOutputParser()
    
    def _generate_batch_tasks(self, story_batch: List[Dict], project_context: Dict, starting_task_id: int) -> List[Dict]:
        """Process multiple stories in a single LLM call for improved performance"""
        
        batch_start_time = datetime.now()
        tech_stack = project_context.get("tech_stack", ["Python", "React", "PostgreSQL"])
        tech_stack_str = ", ".join(tech_stack)
        
        # Format stories concisely for the prompt
        stories_formatted = []
        for story in story_batch:
            story_text = f"{story['id']}: {story['title']} | {story.get('priority', 'medium')} priority"
            stories_formatted.append(story_text)
        
        stories_formatted_str = "\n".join(stories_formatted)
        
        try:
            chain = self.batch_task_prompt | self.llm | self.parser
            
            print(f"[TASK_GEN] Batch processing {len(story_batch)} stories...")
            
            tasks = chain.invoke({
                "stories_formatted": stories_formatted_str,
                "tech_stack": tech_stack_str
            })
            
            # Ensure we have a list
            if not isinstance(tasks, list):
                tasks = [tasks] if tasks else []
            
            validated_tasks = []
            task_id_counter = starting_task_id
            
            for task in tasks:
                if isinstance(task, dict):
                    # Ensure required fields and validate story_id
                    story_id = task.get("story_id", "")
                    if not story_id or not any(s["id"] == story_id for s in story_batch):
                        # Try to infer story_id from task content or assign to first story
                        story_id = story_batch[0]["id"]
                    
                    validated_task = {
                        "id": f"T{task_id_counter:03d}",
                        "story_id": story_id,
                        "title": task.get("title", f"Task for {story_id}"),
                        "description": task.get("description", task.get("title", "")),
                        "category": task.get("category", "backend"),
                        "estimated_hours": min(16, max(4, task.get("estimated_hours", 8))),
                        "priority": task.get("priority", "medium"),
                        "dependencies": task.get("dependencies", []),
                        "acceptance_criteria": task.get("acceptance_criteria", ["Task completed successfully"]),
                        "technical_notes": task.get("technical_notes", "")
                    }
                    validated_tasks.append(validated_task)
                    task_id_counter += 1
            
            elapsed = (datetime.now() - batch_start_time).total_seconds()
            print(f"[TASK_GEN] Batch completed: {len(story_batch)} stories -> {len(validated_tasks)} tasks in {elapsed:.1f}s")
            
            return validated_tasks
            
        except Exception as e:
            print(f"[TASK_GEN_ERROR] Batch processing failed: {e}")
            # Fallback to individual story processing
            return self._create_fallback_tasks_for_batch(story_batch, project_context, starting_task_id)
    
    def _create_fallback_tasks_for_batch(self, story_batch: List[Dict], project_context: Dict, starting_task_id: int) -> List[Dict]:
        """Create fallback tasks for a batch of stories when LLM fails"""
        print(f"[TASK_GEN] Creating fallback tasks for {len(story_batch)} stories")
        
        all_tasks = []
        task_counter = starting_task_id - 1
        
        for story in story_batch:
            fallback_tasks = self._create_fallback_tasks(story, task_counter)
            task_counter += len(fallback_tasks)
            all_tasks.extend(fallback_tasks)
        
        return all_tasks
    
    def _create_fallback_tasks(self, story: Dict, task_counter: int) -> List[Dict]:
        """Create basic fallback tasks when LLM decomposition fails"""
        story_id = story["id"]
        base_title = story["title"].replace("As a ", "").replace(", I want ", " - ").replace(" so that ", " - ")
        
        fallback_tasks = [
            {
                "id": f"T{task_counter + 1:03d}",
                "story_id": story_id,
                "title": f"Backend implementation for {base_title}",
                "description": f"Implement backend logic and API endpoints for {story['description']}",
                "category": "backend",
                "estimated_hours": 12,
                "priority": story.get("priority", "medium"),
                "dependencies": [],
                "acceptance_criteria": story.get("acceptance_criteria", [])[:2],
                "technical_notes": story.get("technical_notes", "")
            },
            {
                "id": f"T{task_counter + 2:03d}",
                "story_id": story_id,
                "title": f"Frontend implementation for {base_title}",
                "description": f"Create UI components and user interface for {story['description']}",
                "category": "frontend",
                "estimated_hours": 10,
                "priority": story.get("priority", "medium"),
                "dependencies": [f"T{task_counter + 1:03d}"],
                "acceptance_criteria": story.get("acceptance_criteria", [])[2:4] if len(story.get("acceptance_criteria", [])) > 2 else ["UI displays correctly", "User interactions work as expected"],
                "technical_notes": "Ensure responsive design and accessibility"
            },
            {
                "id": f"T{task_counter + 3:03d}",
                "story_id": story_id,
                "title": f"Testing for {base_title}",
                "description": f"Write and execute tests for {story['description']}",
                "category": "testing",
                "estimated_hours": 6,
                "priority": "medium",
                "dependencies": [f"T{task_counter + 1:03d}", f"T{task_counter + 2:03d}"],
                "acceptance_criteria": ["Unit tests pass", "Integration tests pass", "Code coverage > 80%"],
                "technical_notes": "Include both unit and integration tests"
            }
        ]
        
        return fallback_tasks
    
    def generate_tasks(self, state: ProjectManagementState) -> ProjectManagementState:
        """Main method to generate tasks from validated user stories using intelligent batching"""
        start_time = datetime.now()
        
        try:
            user_stories = state.get("user_stories", [])
            project_context = state.get("project_context", {})
            
            if not user_stories:
                state["last_error"] = "No user stories available for task generation"
                state["current_phase"] = "error"
                return state
            
            num_stories = len(user_stories)
            print(f"[TASK_GEN] Generating tasks for {num_stories} user stories using batch processing")
            
            # Intelligent batching logic
            if num_stories <= 7:
                # Single batch for small sets
                batch_size = num_stories
                num_batches = 1
                print(f"[TASK_GEN] Using single batch processing for {num_stories} stories")
            else:
                # Multiple batches of 7 stories each for better efficiency
                batch_size = 7
                num_batches = (num_stories + batch_size - 1) // batch_size
                print(f"[TASK_GEN] Processing {num_stories} stories in {num_batches} batches of {batch_size}")
            
            all_tasks = []
            task_id_counter = 1
            
            # Process stories in batches
            for i in range(num_batches):
                batch_start_idx = i * batch_size
                batch_end_idx = min((i + 1) * batch_size, num_stories)
                story_batch = user_stories[batch_start_idx:batch_end_idx]
                
                print(f"[TASK_GEN] Processing batch {i+1}/{num_batches}: stories {batch_start_idx+1}-{batch_end_idx}")
                
                # Generate tasks for this batch
                batch_tasks = self._generate_batch_tasks(story_batch, project_context, task_id_counter)
                
                # Update task counter for next batch
                task_id_counter += len(batch_tasks)
                all_tasks.extend(batch_tasks)
            
            # Ensure all stories have tasks (fallback for missing ones)
            stories_with_tasks = {task["story_id"] for task in all_tasks}
            stories_without_tasks = [story for story in user_stories if story["id"] not in stories_with_tasks]
            
            if stories_without_tasks:
                print(f"[TASK_GEN] Generating fallback tasks for {len(stories_without_tasks)} stories without tasks")
                for story in stories_without_tasks:
                    fallback_tasks = self._create_fallback_tasks(story, task_id_counter - 1)
                    task_id_counter += len(fallback_tasks)
                    all_tasks.extend(fallback_tasks)
            
            # Ensure sequential task IDs
            for i, task in enumerate(all_tasks):
                task["id"] = f"T{i + 1:03d}"
            
            # Update state
            state["tasks"] = all_tasks
            state["current_phase"] = "task_assignment"
            
            processing_time = (datetime.now() - start_time).total_seconds()
            if "processing_time" not in state:
                state["processing_time"] = {}
            state["processing_time"]["task_generation"] = processing_time
            
            # Performance summary
            avg_time_per_story = processing_time / num_stories if num_stories > 0 else 0
            print(f"[TASK_GEN] Generated {len(all_tasks)} tasks for {num_stories} stories in {processing_time:.1f}s")
            print(f"[TASK_GEN] Performance: {avg_time_per_story:.1f}s per story, {num_batches} batches")
            
        except Exception as e:
            state["last_error"] = f"Task generation failed: {str(e)}"
            state["tasks"] = []
            state["current_phase"] = "error"
            print(f"[TASK_GEN_ERROR] {str(e)}")
            import traceback
            traceback.print_exc()
            
        return state

# ==================== WORKFLOW SETUP ====================

def create_story_workflow(gemini_api_key: str = None, max_iterations: int = 3) -> StateGraph:
    """Create the LangGraph workflow for story generation and validation with feedback loop"""
    
    # Initialize agents
    story_agent = MultimodalUserStoryGenerationAgent(gemini_api_key)
    validation_agent = EnhancedUserStoryValidationAgent(gemini_api_key)
    task_agent = TaskGenerationAgent(gemini_api_key)
    
    # Create workflow
    workflow = StateGraph(ProjectManagementState)
    
    def initialize_iteration_tracking(state: ProjectManagementState) -> ProjectManagementState:
        """Initialize iteration tracking in state"""
        if "iteration_count" not in state:
            state["iteration_count"] = 0
        if "max_iterations" not in state:
            state["max_iterations"] = max_iterations
        return state
    
    # Add nodes
    workflow.add_node("initialize", initialize_iteration_tracking)
    workflow.add_node("generate_stories", story_agent.generate_stories)
    workflow.add_node("validate_stories", validation_agent.validate_stories)
    workflow.add_node("generate_tasks", task_agent.generate_tasks)
    
    # Add placeholder nodes
    workflow.add_node("task_creation", lambda x: x)  # Placeholder
    workflow.add_node("requirement_parsing", lambda x: x)  # Placeholder
    workflow.add_node("human_review", lambda x: x)  # Placeholder
    
    # Add edges
    workflow.add_edge("initialize", "generate_stories")
    workflow.add_edge("generate_stories", "validate_stories")
    workflow.add_edge("generate_tasks", "task_creation")
    
    # Enhanced conditional edges with feedback loop
    def determine_next_step(state: ProjectManagementState) -> str:
        validation_status = state.get("validation_status")
        iteration_count = state.get("iteration_count", 0)
        max_iter = state.get("max_iterations", max_iterations)
        validation_score = state.get("validation_score", 0)
        
        print(f"[WORKFLOW] Iteration {iteration_count}, Status: {validation_status}, Score: {validation_score:.1f}")
        
        # Check if approved or good enough
        if validation_status == ValidationStatus.APPROVED.value or validation_score >= 85:
            return "approved"
        
        # If we've reached max iterations but have decent stories, proceed
        if iteration_count >= max_iter:
            if validation_score >= 70:
                print(f"[WORKFLOW] Max iterations reached, score {validation_score:.1f} is acceptable.")
                return "approved"
            else:
                print(f"[WORKFLOW] Max iterations reached, escalating to human review")
                return "max_iterations"
        
        # Continue with revision if possible
        if validation_score >= 30:
            print(f"[WORKFLOW] Score {validation_score:.1f} suggests improvement possible, iterating...")
            return "needs_revision"
        else:
            print(f"[WORKFLOW] Low score {validation_score:.1f}, escalating to human review")
            return "low_quality"
    
    workflow.add_conditional_edges(
        "validate_stories",
        determine_next_step,
        {
            "approved": "generate_tasks",
            "needs_revision": "generate_stories",
            "max_iterations": "human_review",
            "low_quality": "human_review"
        }
    )
    
    # Set entry point
    workflow.set_entry_point("initialize")
    
    return workflow

# ==================== UTILITY FUNCTIONS ====================

def create_multimodal_documentation(
    primary_requirements: str,
    document_path: Optional[str] = None,
    title: str = "Multimodal Project Requirements"
) -> Dict:
    """
    Create multimodal documentation from primary text and optional document file.
    This is the KEY function that formats content correctly for the multimodal agent.
    """
    
    content_parts = []
    
    # Always include primary requirements section
    if primary_requirements.strip():
        content_parts.append("=== PROJECT REQUIREMENTS (TEXT) ===")
        content_parts.append(primary_requirements.strip())
        content_parts.append("")  # Empty line
    
    # Add document content if provided
    if document_path and os.path.exists(document_path):
        try:
            document_content = _extract_text_from_file(document_path)
            if document_content.strip():
                filename = os.path.basename(document_path)
                content_parts.append(f"=== DOCUMENT: {filename} ===")
                content_parts.append(document_content.strip())
        except Exception as e:
            print(f"[WARNING] Failed to load document {document_path}: {e}")
    
    # Combine all content
    combined_content = "\n".join(content_parts)
    
    return {
        "document_type": "Combined Requirements Document",
        "title": title,
        "content": combined_content,
        "sections": {},
        "stakeholders": [],
        "business_goals": [],
        "technical_constraints": [],
        "success_criteria": []
    }

def _extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return _extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return _extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            # Try as plain text
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        raise ValueError(f"Failed to extract text from {file_path}: {e}")

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from PDF file"""
    try:
        import PyPDF2
        
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
        
        content = '\n'.join(text_content)
        return _clean_extracted_text(content)
        
    except ImportError:
        raise ValueError("PyPDF2 library not installed. Run: pip install PyPDF2")
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")

def _extract_text_from_docx(file_path: str) -> str:
    """Extract text content from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        content = '\n'.join(text_content)
        return _clean_extracted_text(content)
        
    except ImportError:
        raise ValueError("python-docx library not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def _clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text from documents"""
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    
    # Remove strange characters that sometimes appear in PDFs
    text = re.sub(r'[^\w\s\.,;:!?\-()[\]{}"\'/\\+=*&%$#@|<>~`]', '', text)
    
    # Fix common PDF extraction issues
    text = text.replace('•', '-')
    text = text.replace('○', '-')
    text = text.replace('▪', '-')
    
    return text.strip()

# ==================== TEST RUNNER ====================

def test_multimodal_workflow(
    primary_requirements: str,
    document_path: Optional[str] = None,
    project_context: Optional[Dict] = None,
    max_iterations: int = 3
) -> Dict:
    """
    Test the multimodal workflow with given requirements and optional document.
    
    Args:
        primary_requirements: Main text requirements (user input)
        document_path: Optional path to supporting document (PDF, DOCX, TXT)
        project_context: Optional project metadata
        max_iterations: Maximum validation iterations
        
    Returns:
        Dict with results including user stories and validation metrics
    """
    
    # Load environment variables
    try:
        from dotenv import load_dotenv
        load_dotenv()
    except ImportError:
        pass
    
    # Get API key
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("GEMINI_API_KEY environment variable not set")
    
    # Create multimodal documentation
    documentation = create_multimodal_documentation(
        primary_requirements=primary_requirements,
        document_path=document_path,
        title="Test Multimodal Requirements"
    )
    
    # Setup project context
    default_context = {
        "industry": "Technology",
        "team_size": 8,
        "tech_stack": ["Python", "React", "PostgreSQL"],
        "timeline": "4 months",
        "budget": "medium"
    }
    if project_context:
        default_context.update(project_context)
    
    # Create initial state
    initial_state = {
        "project_id": f"TEST_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
        "client_requirements": "",  # Empty - using documentation
        "documentation": documentation,
        "project_context": default_context,
        "current_phase": "story_generation",
        "iteration_count": 0,
        "timestamp": datetime.now().isoformat()
    }
    
    # Run workflow
    workflow = create_story_workflow(gemini_api_key, max_iterations)
    app = workflow.compile()
    
    print(f"🚀 Testing multimodal workflow...")
    print(f"📝 Primary requirements: {len(primary_requirements)} chars")
    if document_path:
        print(f"📄 Document: {document_path}")
    
    final_result = None
    iterations = []
    
    for output in app.stream(initial_state):
        for key, value in output.items():
            final_result = value
            
            if key == "generate_stories":
                iteration = value.get("iteration_count", 0)
                story_count = len(value.get('user_stories', []))
                print(f"📝 Iteration {iteration + 1}: Generated {story_count} stories")
            
            elif key == "validate_stories":
                iteration = value.get("iteration_count", 0)
                status = value.get('validation_status', 'unknown')
                score = value.get('validation_score', 0)
                
                iterations.append({
                    "iteration": iteration,
                    "status": status,
                    "score": score,
                    "story_count": len(value.get('user_stories', []))
                })
                
                print(f"✅ Iteration {iteration}: {status} (Score: {score:.1f}/100)")
    
    # Return comprehensive results
    if final_result and final_result.get('user_stories'):
        return {
            "success": True,
            "workflow_results": {
                "iterations": iterations,
                "final_score": final_result.get('validation_score', 0),
                "final_status": final_result.get('validation_status'),
                "total_iterations": len(iterations)
            },
            "user_stories": final_result['user_stories'],
            "tasks": final_result.get('tasks', []),  # ✅ Fix: Extract from final_result
            "multimodal_metadata": final_result.get('multimodal_metadata', {}),
            "processing_time": final_result.get('processing_time', {}),
            "metadata": {
                "generated_at": datetime.now().isoformat(),
                "story_count": len(final_result['user_stories']),
                "validation_score": final_result.get('validation_score', 0)
            }
        }
    else:
        return {
            "success": False,
            "error": final_result.get('last_error', 'Unknown error occurred'),
            "final_state": final_result
        }

# ==================== USAGE EXAMPLES ====================

if __name__ == "__main__":
    
    # Example 1: Text + PDF (Multimodal)
    primary_text = """
        We're building an exclusive online marketplace called 'ArtisanCraft' for high-end, custom-made furniture. Our vision is to connect skilled artisans directly with customers who appreciate unique, handcrafted pieces.

        The experience needs to feel premium and seamless. The most important features are:

        1.  **A beautiful gallery:** Customers need to be able to browse through different furniture pieces made by our artisans. It should be very visual.
        2.  **The customization tool:** This is our main selling point. Customers must be able to select a base furniture design and customize it. For a chair, for example, they should be able to choose the type of wood, the fabric for the upholstery, and even the finish. As they make changes, the price should update in real-time.
        3.  **Artisan Portfolios:** Each artisan needs a profile page where they can showcase their work, share their story, and manage their product listings.
        4.  **A simple, secure checkout:** The payment process has to be straightforward and build trust.

        We're targeting a launch in 6 months, so we need to focus on getting these core features right.
    """
    
    # Path to a sample PDF (create a sample or use None for text-only)
    sample_pdf_path = "requirements-doc.pdf"  # Set to actual PDF path if you have one
    # sample_pdf_path = None  # Use this for text-only testing
    
    try:
        # Run the test
        results = test_multimodal_workflow(
            primary_requirements=primary_text,
            document_path=sample_pdf_path if os.path.exists(sample_pdf_path or "") else None,
            project_context={
                "industry": "E-commerce / Retail Technology",
                "team_size": 6,
                "tech_stack": ["MongoDB", "Express.js", "React", "Node.js"],
                "timeline": "6 months",
                "budget": "High"
            },
            max_iterations=3
        )
        
        if results["success"]:
            print("\n" + "="*80)
            print("🎉 WORKFLOW COMPLETED SUCCESSFULLY")
            print("="*80)
            print(f"Final Score: {results['workflow_results']['final_score']:.1f}/100")
            print(f"Total Iterations: {results['workflow_results']['total_iterations']}")
            print(f"Stories Generated: {results['metadata']['story_count']}")
            
            # Show multimodal metadata
            mm_metadata = results.get("multimodal_metadata", {})
            if mm_metadata:
                print(f"\n📊 Multimodal Analysis:")
                print(f"  - Source Distribution: {mm_metadata.get('source_distribution', {})}")
                print(f"  - Stories with Primary Coverage: {mm_metadata.get('stories_with_primary_coverage', 0)}")
                print(f"  - Stories with Document Coverage: {mm_metadata.get('stories_with_document_coverage', 0)}")
            
            # Show first few stories
            print(f"\n📝 Sample User Stories:")
            for i, story in enumerate(results["user_stories"][:3]):
                print(f"\n{i+1}. {story['id']}: {story['title']}")
                print(f"   Priority: {story['priority']} | Points: {story['estimated_points']}")
                print(f"   Acceptance Criteria: {len(story['acceptance_criteria'])} items")
                
        else:
            print("❌ WORKFLOW FAILED")
            print(f"Error: {results.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ Test failed with error: {e}")
        import traceback
        traceback.print_exc()