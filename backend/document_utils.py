# ==================== UTILITY FUNCTIONS ====================

import os
import re
from typing import Optional, Dict, Any

def create_multimodal_documentation(
    primary_requirements: str,
    document_path: Optional[str] = None,
    title: str = "Multimodal Project Requirements"
) -> Dict:
    """
    Create multimodal documentation from primary text and optional document file.
    This is the KEY function that formats content correctly for the multimodal agent.
    """
    
    content_parts = []
    
    # Always include primary requirements section
    if primary_requirements.strip():
        content_parts.append("=== PROJECT REQUIREMENTS (TEXT) ===")
        content_parts.append(primary_requirements.strip())
        content_parts.append("")  # Empty line
    
    # Add document content if provided
    if document_path and os.path.exists(document_path):
        try:
            document_content = _extract_text_from_file(document_path)
            if document_content.strip():
                filename = os.path.basename(document_path)
                content_parts.append(f"=== DOCUMENT: {filename} ===")
                content_parts.append(document_content.strip())
        except Exception as e:
            print(f"[WARNING] Failed to load document {document_path}: {e}")
    
    # Combine all content
    combined_content = "\n".join(content_parts)
    
    return {
        "document_type": "Combined Requirements Document",
        "title": title,
        "content": combined_content,
        "sections": {},
        "stakeholders": [],
        "business_goals": [],
        "technical_constraints": [],
        "success_criteria": []
    }

def _extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats"""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return _extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return _extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            # Try as plain text
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        raise ValueError(f"Failed to extract text from {file_path}: {e}")

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from PDF file with better handling of spacing issues"""
    
    # Try pdfplumber first (better text extraction)
    try:
        import pdfplumber
        
        print(f"[PDF] Using pdfplumber for extraction: {file_path}")
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    print(f"[WARNING] Error extracting text from page {page_num + 1} with pdfplumber: {e}")
                    continue
        
        if text_content:
            content = '\n\n'.join(text_content)
            cleaned = _clean_extracted_text(content)
            print(f"[PDF] Extracted {len(cleaned)} characters using pdfplumber")
            return cleaned
        else:
            print("[WARNING] pdfplumber returned no content, falling back to PyPDF2")
            
    except ImportError:
        print("[INFO] pdfplumber not installed, falling back to PyPDF2")
    except Exception as e:
        print(f"[WARNING] pdfplumber extraction failed: {e}, falling back to PyPDF2")
        import traceback
        traceback.print_exc()
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        
        print(f"[PDF] Using PyPDF2 for extraction (fallback): {file_path}")
        text_content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    print(f"[WARNING] Error extracting text from page {page_num + 1}: {e}")
                    continue
        
        content = '\n\n'.join(text_content)
        cleaned = _clean_extracted_text(content)
        print(f"[PDF] Extracted {len(cleaned)} characters using PyPDF2")
        return cleaned
        
    except ImportError:
        raise ValueError("PDF extraction libraries not installed. Run: pip install PyPDF2 pdfplumber")
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")

def _extract_text_from_docx(file_path: str) -> str:
    """Extract text content from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        content = '\n'.join(text_content)
        return _clean_extracted_text(content)
        
    except ImportError:
        raise ValueError("python-docx library not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def _clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text from documents"""
    # Remove excessive whitespace while preserving paragraph structure
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple blank lines -> double newline
    text = re.sub(r' +', ' ', text)  # Multiple spaces -> single space
    text = re.sub(r' \n', '\n', text)  # Remove trailing spaces before newlines
    text = re.sub(r'\n ', '\n', text)  # Remove leading spaces after newlines
    
    # Fix common PDF extraction issues with bullet points
    text = text.replace('•', '-')
    text = text.replace('○', '-')
    text = text.replace('▪', '-')
    # NOTE: Removed empty string replacement that was causing corruption
    
    # Remove null bytes and other control characters that may cause issues
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
    
    # Fix common ligatures and special characters from PDFs
    text = text.replace('ﬁ', 'fi')
    text = text.replace('ﬂ', 'fl')
    text = text.replace('ﬀ', 'ff')
    text = text.replace('ﬃ', 'ffi')
    text = text.replace('ﬄ', 'ffl')
    text = text.replace('—', '-')
    text = text.replace('–', '-')
    text = text.replace(''', "'")
    text = text.replace(''', "'")
    text = text.replace('"', '"')
    text = text.replace('"', '"')
    text = text.replace('…', '...')
    
    # Fix hyphenated line breaks (common in PDFs)
    # "word-\nword" -> "word-word" (keep hyphen)
    # "word\nword" -> "word word" (add space for split words)
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1-\2', text)  # Keep hyphens
    text = re.sub(r'(\w)\s*\n\s*(\w)', r'\1 \2', text)   # Add space for split words
    
    # Remove any remaining non-printable characters but keep all text and common punctuation
    # This is much less aggressive than before
    text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
    
    return text.strip()